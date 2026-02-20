#!/usr/bin/env python3
"""
Document Reader System - Multi-format document processing
Handles PDF, DOCX, images with OCR, and more
"""

import os
import subprocess
import json
from pathlib import Path

class DocumentReader:
    """
    Universal document reader with OCR capabilities
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.read_pdf,
            '.docx': self.read_docx,
            '.doc': self.read_docx,
            '.txt': self.read_text,
            '.md': self.read_text,
            '.png': self.read_image,
            '.jpg': self.read_image,
            '.jpeg': self.read_image,
            '.webp': self.read_image,
        }
        self.temp_dir = "/tmp/doc_reader"
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def read_file(self, file_path: str) -> dict:
        """Read any supported file format"""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext not in self.supported_formats:
            return {
                'success': False,
                'error': f'Unsupported format: {ext}',
                'supported': list(self.supported_formats.keys())
            }
        
        try:
            reader = self.supported_formats[ext]
            content = reader(file_path)
            return {
                'success': True,
                'format': ext,
                'content': content,
                'file': file_path
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'file': file_path
            }
    
    def read_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdftotext or pdfplumber"""
        # Try pdftotext first (faster)
        result = subprocess.run(
            ['pdftotext', file_path, '-'],
            capture_output=True,
            text=True
        )
        
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
        
        # Fallback to Python library
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                return text
        except ImportError:
            return "PDF extraction requires: sudo apt-get install poppler-utils"
    
    def read_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    text_parts.append(row_text)
            
            return '\n'.join(text_parts)
        except ImportError:
            return "DOCX extraction requires: pip install python-docx"
    
    def read_text(self, file_path: str) -> str:
        """Read plain text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def read_image(self, file_path: str) -> str:
        """Extract text from images using OCR"""
        try:
            from PIL import Image
            import pytesseract
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except ImportError:
            return "OCR requires: pip install pytesseract pillow && apt-get install tesseract-ocr"
        except Exception as e:
            return f"OCR failed: {str(e)}"
    
    def batch_process(self, directory: str) -> list:
        """Process all supported files in directory"""
        results = []
        for file_path in Path(directory).glob('*'):
            if file_path.suffix.lower() in self.supported_formats:
                result = self.read_file(str(file_path))
                results.append(result)
        return results

def install_dependencies():
    """Install required packages"""
    packages = [
        'python-docx',
        'pdfplumber',
        'Pillow',
        'pytesseract',
    ]
    
    print("Installing dependencies...")
    for pkg in packages:
        subprocess.run(['pip', 'install', '-q', pkg])
    
    print("✅ Dependencies installed")
    print("\nSystem packages needed:")
    print("  sudo apt-get install poppler-utils tesseract-ocr")

def main():
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python document_reader.py <file_path>")
        print("\nSupported formats:")
        print("  PDF (.pdf)")
        print("  Word (.docx, .doc)")
        print("  Images (.png, .jpg, .jpeg, .webp)")
        print("  Text (.txt, .md)")
        return
    
    file_path = sys.argv[1]
    
    if file_path == '--install':
        install_dependencies()
        return
    
    reader = DocumentReader()
    result = reader.read_file(file_path)
    
    if result['success']:
        print(f"✅ Successfully read: {result['file']}")
        print(f"Format: {result['format']}")
        print("\n" + "="*50)
        print(result['content'][:2000])  # First 2000 chars
        if len(result['content']) > 2000:
            print("\n... [truncated]")
    else:
        print(f"❌ Error: {result['error']}")
        if 'supported' in result:
            print(f"Supported formats: {', '.join(result['supported'])}")

if __name__ == "__main__":
    main()
